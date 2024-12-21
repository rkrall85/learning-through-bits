


# need to install python docx module
# pip install python-docx




# Document object contains paragraph objects.
# Paragraph objects contain run objects

import docx
import os

os.chdir('C:\\Users\\rkrall\\github\\AutomateTheBoringStuffWithPython\\14_ExcelWordPDF')
d = docx.Document('45_xdemo.docx')

#print(d.paragraphs)


#print(d.paragraphs[1].text) #output the paragraph text

p = d.paragraphs[1]
#print(p.runs) # a new run whenever the style changes (run is a style change)
#print(p.runs[0].text)
#print(p.runs[1].text)
#print(p.runs[2].text)
#print(p.runs[3].text)

#output true or false if its bold/italic/underline text
#print(p.runs[1].bold)
#print(p.runs[1].italic)
#print(p.runs[1].underline)

#change run style
#print(p.runs[1].underline)
p.runs[1].underline = True
p.runs[1].text = 'italic and underlined'
#print(p.runs[1].underline)

p.style = 'Title' #change style to title style



d.save('45_xdemo2.docx')


d = docx.Document()
d.add_paragraph('Hello thi is a paragraph')
d.add_paragraph('This is another paragraph')
p = d.paragraphs[0]
p.add_run(' This is a new run.')
p.runs[1].bold = True

d.save('45_xdemo3.docx')
